# By Constantine Sidorov, 2026

from __future__ import annotations

import datetime as dt
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import openpyxl
from docxtpl import DocxTemplate
from docx import Document

from excel_headers import read_headers

@dataclass(frozen=True)
class ExcelRow:
    no: str | None
    name: str | None
    kind: str | None
    qty: str | None
    inv: str | None
    serial: str | None
    year: str | None
    defect_mon: str | None
    defect_sys: str | None


@dataclass(frozen=True)
class ActData:
    inv_no: str
    kit_name: str
    def_mon: str
    def_sys: str
    sys_line: str
    mon_line: str
    sys_serial_line: str
    mon_serial_line: str


def _find_single(root: Path, pattern: str) -> Path:
    matches = sorted(root.glob(pattern))
    if not matches:
        raise FileNotFoundError(f"Не найдено по шаблону: {pattern}")
    if len(matches) > 1:
        raise RuntimeError(f"Ожидался один файл по {pattern}, найдено: {matches}")
    return matches[0]


def _norm(v: Any) -> str | None:
    if v is None:
        return None
    s = str(v).strip()
    return s if s else None


def _parse_inv_base(inv: str) -> str:
    return inv.split("/", 1)[0].strip()


def _is_base_row(inv: str) -> bool:
    return "/" not in inv


def _guess_part_role(kind: str | None, name: str | None) -> str | None:
    text = " ".join([k for k in [kind, name] if k]).lower()
    if not text:
        return None
    if any(w in text for w in ["монитор", "display", "viewsonic", "proview"]):
        return "monitor"
    if any(w in text for w in ["систем", "pc", "unit", "aquarius"]):
        return "system"
    return None


def read_sheet_rows(xlsx_path: Path) -> list[ExcelRow]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[wb.sheetnames[0]]

    out: list[ExcelRow] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        # A..I
        no, name, kind, qty, inv, serial, year, defect_mon, defect_sys = (list(row) + [None] * 9)[:9]
        out.append(
            ExcelRow(
                no=_norm(no),
                name=_norm(name),
                kind=_norm(kind),
                qty=_norm(qty),
                inv=_norm(inv),
                serial=_norm(serial),
                year=_norm(year),
                defect_mon=_norm(defect_mon),
                defect_sys=_norm(defect_sys),
            )
        )
    return out


def build_acts(rows: list[ExcelRow]) -> list[ActData]:
    by_base: dict[str, list[ExcelRow]] = {}
    for r in rows:
        if not r.inv:
            continue
        base = _parse_inv_base(r.inv)
        if not re.fullmatch(r"\d{6,}", base):
            continue
        by_base.setdefault(base, []).append(r)

    acts: list[ActData] = []
    for inv_no, group in sorted(by_base.items()):
        base_rows = [r for r in group if r.inv == inv_no]
        if not base_rows:
            continue
        base = base_rows[0]

        kit_name = base.name or ""
        def_mon = base.defect_mon or ""
        def_sys = base.defect_sys or ""

        sys_name = ""
        mon_name = ""
        sys_serial = ""
        mon_serial = ""

        part_rows = [r for r in group if r.inv and not _is_base_row(r.inv)]
        for pr in part_rows:
            role = _guess_part_role(pr.kind, pr.name)
            if role == "system" and not sys_name:
                sys_name = pr.name or ""
                sys_serial = pr.serial or ""
            elif role == "monitor" and not mon_name:
                mon_name = pr.name or ""
                mon_serial = pr.serial or ""

        if not all([kit_name, sys_name, mon_name]):
            # минимально строгая валидация: пропускаем кривые/неполные группы
            continue

        sys_line = f"Системный блок: {sys_name}"
        mon_line = f"Монитор: {mon_name}"
        sys_serial_line = f"Системный блок: {sys_serial}" if sys_serial else "Системный блок: "
        mon_serial_line = f"Монитор: {mon_serial}" if mon_serial else "Монитор: "

        acts.append(
            ActData(
                inv_no=inv_no,
                kit_name=kit_name,
                def_mon=def_mon,
                def_sys=def_sys,
                sys_line=sys_line,
                mon_line=mon_line,
                sys_serial_line=sys_serial_line,
                mon_serial_line=mon_serial_line,
            )
        )
    return acts


def _act_date(today: dt.date) -> str:
    months = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }
    day = f"{today.day:02d}"
    return f"«{day}»  {months[today.month]}  {today.year}   г."


def _read_dop_json(dop_path: Path | None) -> dict[str, str]:
    """
    DOP.JSON:
      {
        "ACT_DATE": "...",
        "BASIS": "...",
        "MOL": "..."
      }
    """
    defaults = {
        "ACT_DATE": "«26»  февраля  2026   г.",
        "BASIS": "служебная записка от 05.02.2026 №4.3-00444-СЗ",
        "MOL": "Заместитель начальника отдела, отдел подготовки сведений, Романович Елена Алексеевна",
    }
    if dop_path is None or not dop_path.exists():
        return defaults
    try:
        data = json.loads(dop_path.read_text(encoding="utf-8"))
    except Exception:
        return defaults

    out: dict[str, str] = {}
    for k, v in defaults.items():
        vv = data.get(k)
        out[k] = str(vv).strip() if vv is not None and str(vv).strip() else v
    return out


def _clear_highlight(doc_path: Path) -> None:
    doc = Document(str(doc_path))

    def clear_paragraph(p) -> None:
        for run in p.runs:
            if run.font.highlight_color is not None:
                run.font.highlight_color = None

    for p in doc.paragraphs:
        clear_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    clear_paragraph(p)

    doc.save(str(doc_path))


def render_all(
    *,
    template_docx: Path,
    excel_xlsx: Path,
    acts: list[ActData],
    out_dir: Path,
    dop_json: Path | None,
) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)

    h = read_headers(excel_xlsx)
    dop = _read_dop_json(dop_json)

    for act in acts:
        tpl = DocxTemplate(str(template_docx))
        ctx = {
            h.inv_no: act.inv_no,
            h.name: act.kit_name,
            f"{h.name}_sys": act.sys_line.replace("Системный блок: ", "", 1),
            f"{h.name}_mon": act.mon_line.replace("Монитор: ", "", 1),
            f"{h.serial}_sys": act.sys_serial_line.replace("Системный блок: ", "", 1).strip(),
            f"{h.serial}_mon": act.mon_serial_line.replace("Монитор: ", "", 1).strip(),
            h.defect_mon: act.def_mon,
            h.defect_sys: act.def_sys,
            "ACT_DATE": dop["ACT_DATE"] or _act_date(dt.date.today()),
            "BASIS": dop["BASIS"],
            "MOL": dop["MOL"],
        }
        tpl.render(ctx)
        dst = out_dir / f"{act.inv_no}.docx"
        tpl.save(str(dst))
        _clear_highlight(dst)


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    xlsx = _find_single(repo_root, "*.xlsx")
    template = repo_root / "template_act.docx"
    if not template.exists():
        raise FileNotFoundError(
            "Не найден template_act.docx. Сначала запусти tools/make_template.py"
        )

    rows = read_sheet_rows(xlsx)
    acts = build_acts(rows)
    if not acts:
        raise RuntimeError("Не удалось собрать ни одного акта из Excel (проверь данные).")

    render_all(
        template_docx=template,
        excel_xlsx=xlsx,
        acts=acts,
        out_dir=repo_root / "out",
        dop_json=repo_root / "DOP.JSON",
    )
    print(f"OK: сгенерировано файлов: {len(acts)} -> {repo_root / 'out'}")


if __name__ == "__main__":
    main()

