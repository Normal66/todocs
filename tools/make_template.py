# By Constantine Sidorov, 2026

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from excel_headers import HeaderMap, read_headers


@dataclass(frozen=True)
class TemplatePaths:
    """
    source_docx: исходный Word-образец. Поля для подстановки подсвечены жёлтым.
    output_docx: сюда сохраняется Word-шаблон с плейсхолдерами {{...}}.
    excel_xlsx: Excel, чтобы назвать плейсхолдеры по заголовкам столбцов.
    """

    source_docx: Path
    output_docx: Path
    excel_xlsx: Path | None = None


def _find_single(root: Path, pattern: str) -> Path:
    matches = sorted(root.glob(pattern))
    if not matches:
        raise FileNotFoundError(f"Не найдено по шаблону: {pattern}")
    if len(matches) > 1:
        raise RuntimeError(f"Ожидался один файл по {pattern}, найдено: {matches}")
    return matches[0]


def _default_paths(repo_root: Path) -> TemplatePaths:
    docxs = sorted(repo_root.glob("*.docx"))
    if not docxs:
        raise FileNotFoundError("В папке проекта не найдено ни одного .docx")
    for p in docxs:
        if p.name == "Образец акт.docx":
            source = p
            break
    else:
        source = docxs[0]
        for p in docxs:
            if not p.stem[:1].isdigit():
                source = p
                break

    xlsx = _find_single(repo_root, "*.xlsx")
    return TemplatePaths(
        source_docx=source,
        output_docx=repo_root / "template_act.docx",
        excel_xlsx=xlsx,
    )


def _replace_all(text: str, replacements: dict[str, str]) -> str:
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _iter_paragraphs(doc: Document):
    """Обходит все параграфы документа: обычные + внутри таблиц."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _yellow_groups(paragraph) -> list[tuple[int, int]]:
    """
    Возвращает группы runs с жёлтой подсветкой.
    Группа = подряд идущие runs, которые выделены жёлтым и содержат текст.
    """
    runs = paragraph.runs
    groups: list[tuple[int, int]] = []
    start: int | None = None

    for i, r in enumerate(runs):
        is_yellow = r.font.highlight_color == WD_COLOR_INDEX.YELLOW and (r.text or "").strip()
        if is_yellow and start is None:
            start = i
        if (not is_yellow) and start is not None:
            groups.append((start, i - 1))
            start = None
    if start is not None:
        groups.append((start, len(runs) - 1))
    return groups


def _placeholder_for(paragraph_text: str, group_text: str, where: str, h: HeaderMap) -> str:
    """
    По тексту строки и месту (таблица/текст) решаем, какой плейсхолдер поставить.
    Имена плейсхолдеров берём из заголовков Excel (HeaderMap).
    """
    if "Заключение №" in paragraph_text:
        return f"{{{{{h.inv_no}}}}}"
    if paragraph_text.strip().startswith("г.") and "Екатеринбург" in paragraph_text:
        return "{{ACT_DATE}}"
    if "Основание проведения экспертизы" in paragraph_text:
        return "{{BASIS}}"
    if "материально ответственным лицом" in paragraph_text:
        return "{{MOL}}"
    if paragraph_text.strip().startswith("Монитор:"):
        if "Вышла" in paragraph_text or "из строя" in paragraph_text:
            return f"{{{{{h.defect_mon}}}}}"
    if paragraph_text.strip().startswith("Системный блок:"):
        if "Вышла" in paragraph_text or "из строя" in paragraph_text:
            return f"{{{{{h.defect_sys}}}}}"
    if where.startswith("T") and "Автоматизированное рабочее место" in paragraph_text:
        return f"{{{{{h.name}}}}}"
    if where.startswith("T") and paragraph_text.strip().startswith("Системный блок:"):
        if any(ch.isdigit() for ch in paragraph_text) and "," not in paragraph_text:
            return f"Системный блок: {{{{{h.serial}_sys}}}}"
        return f"Системный блок: {{{{{h.name}_sys}}}}"
    if where.startswith("T") and paragraph_text.strip().startswith("Монитор:"):
        if any(ch.isdigit() for ch in paragraph_text) and "," not in paragraph_text:
            return f"Монитор: {{{{{h.serial}_mon}}}}"
        return f"Монитор: {{{{{h.name}_mon}}}}"
    if where.startswith("T"):
        if any(ch.isdigit() for ch in group_text) and "/" not in group_text:
            if len("".join(ch for ch in group_text if ch.isdigit())) >= 6:
                return f"{{{{{h.inv_no}}}}}"

    return "{{FIELD}}"


def make_template(paths: TemplatePaths) -> None:
    """
    Берёт Word-образец, находит весь жёлтый текст и превращает его в плейсхолдеры.
    """
    src = paths.source_docx
    dst = paths.output_docx

    doc = Document(str(src))
    if paths.excel_xlsx is None:
        raise ValueError("Нужен путь к Excel-файлу (excel_xlsx).")
    h = read_headers(paths.excel_xlsx)

    field_counter = 1
    for p in _iter_paragraphs(doc):
        groups = _yellow_groups(p)
        if not groups:
            continue

        # python-docx не даёт “адрес” параграфа (таблица/не таблица) в публичном API.
        # Здесь безопасно используем внутреннее поле, чтобы отличить таблицу от текста.
        where = "T" if p._p.getparent().tag.endswith("tc") else "P"  # noqa: SLF001

        paragraph_text = p.text
        squash_all_yellow_to_one: str | None = None
        if paragraph_text.strip().startswith("г.") and "Екатеринбург" in paragraph_text:
            squash_all_yellow_to_one = "{{ACT_DATE}}"
        if "Автоматизированное рабочее место" in paragraph_text:
            squash_all_yellow_to_one = f"{{{{{h.name}}}}}"

        if squash_all_yellow_to_one is not None:
            first_start, first_end = groups[0]
            p.runs[first_start].text = squash_all_yellow_to_one
            for start, end in groups:
                for i in range(start, end + 1):
                    if i == first_start:
                        continue
                    p.runs[i].text = ""
            continue

        for start, end in groups:
            runs = p.runs
            group_text = "".join((runs[i].text or "") for i in range(start, end + 1))

            ph = _placeholder_for(paragraph_text, group_text, where, h)
            if ph == "{{FIELD}}":
                ph = f"{{{{FIELD_{field_counter:03d}}}}}"
                field_counter += 1

            runs[start].text = ph
            for i in range(start + 1, end + 1):
                runs[i].text = ""

    doc.save(str(dst))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    paths = _default_paths(repo_root)
    make_template(paths)
    print(f"OK: шаблон сохранён: {paths.output_docx}")
    print("Проверь файл: жёлтые поля должны стать плейсхолдерами вида {{...}}.")


if __name__ == "__main__":
    main()

